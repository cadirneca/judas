######
# @author:      Ana Nieto
# @country:     Spain
# @website:     https://www.linkedin.com/in/ananietojimenez/
######
#
# Refs docx: https://python-docx.readthedocs.io/en/latest/
#
import xlsxwriter
from docx import Document  #pip3 install python-docx
from docx.shared import Inches
from shutil import copyfile
import fileinput
from PIL import Image, ImageDraw
from email.mime.text import MIMEText
import smtplib, ssl
import mechanize
import random


#--------------------------------------------------------
#   AUXILIARY METHODS
#--------------------------------------------------------
def welcomeStudent(title='SERA Exercise', student='Nieto', evaluation='How this exercise is evaluated'):
    """
    :param title: title of the activity.
    :param student: name of the student.
    :param evaluation: evaluation criteria for this activity.
    :return: string with these fields
    """
    return '%s\nStudent:%s\nEvaluation:%s' % (title, student, evaluation)


def sayHello(document, student, text):
    """
    Includes a 'hello' string for the student in a document.
    :param document: document where the welcome string will be included.
    :param student: name of the student.
    :param text: additional text for the welcome.
    :return: this method will modify the 'document' provided.
    """

    p = document.add_paragraph('Dear ')
    p.add_run(student).bold = True
    p.add_run(' ')
    p.add_run(text).italic = True


def replaceSentence(srcFile, dstFile, targetsentence, newSentence):
    """
    Creates a new file that is a copy of the given file but with the target string replaced.
    :param srcFile: source file to be copied.
    :param dstFile: destination file with the old sentence replaced by the new one.
    :param targetsentence: target sentence to be replaced.
    :param newSentence: new sentence.
    :return: This file creates a new file with a replacement of strings.
    """

    # Create a copy of the base file:
    copyfile(srcFile, dstFile) #copy2 also copies metadata

    # Open file and replace string for student:
    with fileinput.FileInput(dstFile, inplace=True) as file: #, backup='.bak') as file:
        for line in file:
            print(line.replace(targetsentence, newSentence), end='')


def checkEmailServer():
    """
    Checks if fakeemail is running and start the service if isn't
    :return: True when the service is available
    """

#--------------------------------------------------------
#   MAIN METHODS TO GENERATE FILES
#--------------------------------------------------------
def giveMeHelloFile(student, name='demo', ext='.py', ext2=None, path='resources/code/'):
    """
    Generates a new file 'helloworld' using an existing one.
    :param student: Name of the student.
    :param name: name of the file generated (without the extension).
    :param ext: extension of the file.
    :param ext2: None (by default). Alternative extension to be shown to the user.
    :return: This method generates a python file.
    """
    txt = ext
    if ext2 is not None: txt = ext2
    hellomsg = 'Dear %s this is your custom %s file.' % (student, txt)

    # Create a copy of the base file:
    replaceSentence('%shelloworld%s' % (path, ext), name + ext, 'Hello World!', hellomsg)

def giveMeADocx(title='SERA Exercise', student='Nieto', text='SERA exercise description',
                    name='demo.docx', evaluation='How this exercise is evaluated', picture=None):
    """
    Creates a .docx with the following info:
    :param title: title of the document (SERA exercise).
    :param student: name of the student (to say ''hello'').
    :param text: SERA exercise description.
    :param name: name of the .docx.
    :param evaluation: description about how the exercise will be evaluated.
    :return: This method creates a .docx
    """
    document = Document()

    document.add_heading(title, 0)
    sayHello(document, student, 'this is your custom exercise.')

    document.add_paragraph(text)

    document.add_heading('Evaluation', level=1)
    document.add_paragraph(evaluation, style='Intense Quote')

    if picture is not None:
        document.add_picture(picture, width=Inches(3))

    document.add_page_break()
    document.save(name)

def giveMeAXlsx(title='SERA Exercise', student='Nieto', text='SERA exercise description',
                    name='demo.xlsx', evaluation='How this exercise is evaluated', picture=None):
    # Ref: https://xlsxwriter.readthedocs.io/example_demo.html#ex-demo
    # Create an new Excel file and add a worksheet.
    workbook = xlsxwriter.Workbook(name)
    worksheet = workbook.add_worksheet()

    # Widen the first column to make the text clearer.
    worksheet.set_column('A:A', 20)

    # Add a bold format to use to highlight cells.
    bold = workbook.add_format({'bold': True})

    # Write some simple text.
    worksheet.write('A1', title)

    # Text with formatting.
    worksheet.write('A2', student, bold)

    # Write some numbers, with row/column notation.
    worksheet.write(2, 0, text)
    worksheet.write(3, 0, evaluation)

    # Insert an image.
    if picture is not None: worksheet.insert_image('B5', picture)

    workbook.close()



def giveMePicture(title='SERA Exercise', student='Nieto', name='demo', ext = '.png',
                  evaluation='How this exercise is evaluated'):

    img = Image.new('RGB', (300, 300), color=(73, 109, 137))

    d = ImageDraw.Draw(img)
    d.text((10, 10), welcomeStudent(title, student, evaluation), fill=(255, 255, 0))

    img.save(name + ext)


#--------------------------------------------------------
#   EMAIL
#--------------------------------------------------------
# To email spoofer download and install this: https://github.com/mikechabot/smtp-email-spoofer-py
def getFakeEmail():
    # Returns a fake email
    extension = ['gmail.com', 'fakemail.com', 'serasystem.com', 'lcc.uma.es', 'demonios.es', 'demonios.com',
                 'alumnos.com', 'totalmentefalso.com', 'hotmail.com', 'muyfalso.com']
    user = ['serasystem', 'seraforense', 'buho', 'caracola', 'paquito', 'chocolatero', 'babyshark', 'tiroriro',
            'mamashark', 'boqueron', 'malaga', 'fakemail', 'peligro', 'acierto', 'falsamente', 'despacito',
            'poo', 'kk', 'muuu']
    return "%s@%s" % (random.choice(user), random.choice(extension))

def generaHTLM_email(pathfileinput, pathfileoutput, targetsentence, newSentence):

    replaceSentence(pathfileinput, pathfileoutput, targetsentence, newSentence)

    return pathfileoutput


def sendEmail(sender_email, password, dest_email, subject='SERA-email', message='This is a hello message'):
    """
    Sends email using a smtp server ** running in the machine **
    :param sender_email: email address for the origin of the message (must be registered in the server).
    :param password: password of the sender (this email will be legitime, not fake).
    :param dest_email: email address for the reception of the message.
    :param message: text to be sent in the body of the email.
    :return: True if all was fine, false in other case.
    """
    smtp_server = "localhost" #e.g. "smtp.gmail.com"
    port = 25 #1025

    # Create a secure SSL context
    context = ssl.create_default_context()

    # Prepare email:
    msg = MIMEText(message) # content of the body
    msg['Subject'] = subject # subject of the message
    EMAIL_SPACE = ", "
    msg['To'] = EMAIL_SPACE.join(dest_email)
    msg['From'] = sender_email


    # Try to log in to server and send email
    try:
        server = smtplib.SMTP(smtp_server, port)
        server.ehlo()  # Can be omitted
        server.starttls(context=context)  # Secure the connection
        #server.ehlo()  # Can be omitted
        server.login(sender_email, password)

        # Send email
        server.sendmail(sender_email, dest_email, msg.as_string())

    except Exception as e:
        # Print any error messages to stdout
        print(e)
        return False

    finally:
        if server is not None:
            server.quit()

    return True

def sendEmail_anonymous(sender_email,subject='SERA-Message',message='This is a hello message from SERA'):
    """
    Sends an anonymous email using mailtrap.io, the 'sender' is predefined by the tool...
    This method is not fast... it can take around 6 hours in deliver the message to the receipt.
    :param sender_email: destination of the email (target).
    :param subject: subject in the email.
    :param text: text (body) in the email.
    :return: True if the message was delivered, false in other case.
    """

    br = mechanize.Browser()

    url = "http://anonymouse.org/anonemail.html"
    headers = "Mozilla/4.0 (compatible; MSIE 5.0; AOL 4.0; Windows 95; c_athome)"
    br.addheaders = [('User-agent', headers)]
    br.open(url)
    br.set_handle_equiv(True)
    br.set_handle_gzip(True)
    br.set_handle_redirect(True)
    br.set_handle_referer(True)
    br.set_handle_robots(False)
    br.set_debug_http(False)
    br.set_debug_redirects(False)
    # br.set_proxies({"http": proxy})

    br.select_form(nr=0)

    br.form['to'] = sender_email
    br.form['subject'] = subject
    br.form['text'] = message

    result = br.submit()

    response = br.response().read()

    if "The e-mail has been sent anonymously!" in response:
        return True
    else:
        return False






